from tkinter import filedialog
from symptoms_db import SymptomsDB
import os
from docx import Document
from docx.shared import Pt
import openpyxl


class DataManager:
    """Manage gathering data and saving to excel and word"""
    def __init__(self):
        self.dir_path = None
        self.diary_path = None
        self.symptom_path = None

        self.log_data = {}
        self.log_detail_data = {}


    # CHANGE THIS TO GET DIR PATH FROM USER
    def create_export_dir(self):
        """Get export location from user"""
        self.dir_path = filedialog.asksaveasfilename()
        os.mkdir(self.dir_path)
        self.diary_path = os.path.abspath(os.path.join(self.dir_path, 'diary.docx'))
        self.symptom_path = os.path.abspath(os.path.join(self.dir_path, 'symptoms.xlsx'))


    def get_current_data_to_save(self):
        db = SymptomsDB()
        try:
            logs = db.get_logs_in_reverse_date_order()
            for log in logs:
                self.log_data[log[0]] = [log[1], log[3]]
            self.save_diary_data()
            self.save_symptoms_data()
        except Exception as e:
            print("Error: ", e)


    def save_diary_data(self):
        """Save diary data to word"""
        try:
            if os.path.exists(self.diary_path):
                os.remove(self.diary_path)
            diary = Document()
            diary.save(self.diary_path)
            for k, v in self.log_data.items():
                if v[1] != "":
                    diary.add_heading(f"~ Date: {v[0]} ~\n", level=1)
                    paragraph = diary.add_paragraph()
                    run = paragraph.add_run(f"\t{v[1]}")
                    run.font_size = Pt(12)
            diary.save(self.diary_path)
        except Exception as e:
            print("Error: ", e)


    def save_symptoms_data(self):
        """Save symptoms data to excel"""
        db = SymptomsDB()
        for key in self.log_data.keys():
            try:
                data = db.get_log_details_by_id(key)
                for log in data:
                    if log[4] == 'scale':
                        self.log_detail_data[log[0]] = [{
                            "date": self.log_data[key][0],
                            "name": log[2],
                            "level": f"{log[3]}/5"}
                        ]
                    else:
                        self.log_detail_data[log[0]] = [{
                            "date": self.log_data[key][0],
                            "name": log[2],
                            "level": log[3]}
                        ]

            except Exception as e:
                print("Error: ", e)

        symptoms_file = openpyxl.Workbook()
        sheet = symptoms_file.active
        headers = ['Date', 'Symptom', 'Severity']
        for col_num, header in enumerate(headers, start=1):
            sheet.cell(row=1, column=col_num, value=header)
        sheet.column_dimensions["A"].width = 20
        sheet.column_dimensions["B"].width = 20
        sheet.column_dimensions["C"].width = 20

        for i in self.log_detail_data.values():
            sheet.append([i[0]['date'], i[0]['name'], i[0]['level']])

        symptoms_file.save(self.symptom_path)


if __name__ == "__main__":
    test = DataManager()
    test.create_dir_if_not_exists()
    test.get_current_data_to_save()
    test.save_symptoms_data()